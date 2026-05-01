from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re, datetime

def set_yu_gothic(doc):
    style = doc.styles['Normal']
    style.font.name = 'Yu Gothic'
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), 'Yu Gothic')

def safe_filename(title, maxlen=150):
    s = re.sub(r'[<>:"/\\|?*]', '_', title)
    s = s.rstrip(' .')
    return s[:maxlen] + '.docx'

def add_table_row(table, label, value):
    row = table.add_row()
    row.cells[0].text = label
    row.cells[1].text = value

doc = Document()
set_yu_gothic(doc)

title_en = "PHEIGES: all-cell-free phage synthesis and selection from engineered genomes"
title_ja = "PHEIGES：操作されたゲノムからの全無細胞ファージ合成と選択"

doc.add_heading(title_en, level=1)

doc.add_heading("日本語訳タイトル", level=2)
doc.add_paragraph(title_ja)

doc.add_heading("論文情報", level=2)
table = doc.add_table(rows=1, cols=2)
table.style = 'Table Grid'
table.rows[0].cells[0].text = "項目"
table.rows[0].cells[1].text = "内容"
add_table_row(table, "著者", "Levrier A, Karpathakis I, Nash B, Bowden SD, Lindner AB, Noireaux V")
add_table_row(table, "誌名", "Nature Communications")
add_table_row(table, "出版年", "2024")
add_table_row(table, "巻号頁", "15: 2223")
add_table_row(table, "DOI", "10.1038/s41467-024-46585-1")
add_table_row(table, "PMID", "38472230")

doc.add_heading("背景・課題", level=2)
doc.add_paragraph(
    "ファージ工学は従来、宿主菌への形質転換やファージ感染を前提としており、宿主域の制限・高 G+C 含量・必須遺伝子の同定困難さが設計の障壁となっていた。"
    "セルフリー転写翻訳 (TXTL) システムは試験管内でのタンパク質発現を可能にするが、完全な感染性ファージを迅速に合成・選択するワークフローは確立されていなかった。"
    "特に、変異ライブラリーから目的特性を持つファージを直接スクリーニングできる全無細胞プラットフォームが求められていた。"
)

doc.add_heading("手法", level=2)
doc.add_paragraph(
    "PHEIGES (PHage Engineering by In vitro Gene Expression and Selection) と命名された本手法は、T7 ファージゲノムを 12 kbp 以下の PCR 増幅フラグメントに分割し、エクソヌクレアーゼを用いたオーバーラップ末端分解で in vitro アセンブリを行う。"
    "アセンブルされたゲノムを大腸菌由来の TXTL (cell-free transcription-translation) 反応液に直接添加し、室温での 1 バッチ反応でファージ粒子を生産した。"
    "生産されたファージを宿主大腸菌に感染させて選択プレートでスクリーニングし、目的の表現型 (宿主域変異、蛍光タンパク質挿入等) を持つ変異体を単離した。"
    "テール・ファイバー変異ライブラリーを用いたスクリーニングにより、平滑型リポ多糖 (rough LPS) に感染できる T7 変異体を de novo 合成・選択した。"
    "ゲノムの 10% 短縮版および蛍光遺伝子統合版も同一プロセスで作製し、プラットフォームの汎用性を実証した。"
)

doc.add_heading("結果", level=2)
doc.add_paragraph(
    "TXTL 反応液 1 ml あたり最大 10¹¹ PFU の感染性 T7 ファージを 1 日以内に生産できることを確認した。"
    "テール・ファイバー変異ライブラリー (各位置にランダム変異を導入) から、rough LPS を持つ大腸菌株 (通常 T7 が感染できない) に感染可能なファージ変異体を 1 ラウンドのスクリーニングで取得した。"
    "GFP 挿入ゲノムから蛍光ファージの合成に成功し、蛍光強度は野生型と同等の力価を示した。"
    "ゲノムを 10% 短縮 (約 4 kbp 削除) した変異体も感染性を維持し、ミニマルゲノム工学の可能性を示した。"
    "全工程が試験管内で完結するため、宿主菌への依存がなく、設計から感染性ファージ取得まで 24〜48 時間で可能であった。"
    "アセンブリ効率は Gibson Assembly 法と同等以上で、特定の配列バイアスなしに多様な変異体ライブラリーを作製できた。"
)

doc.add_heading("考察・新規性", level=2)
doc.add_paragraph(
    "本研究の最大の新規性は、生きた宿主細胞を一切使用せずにファージを合成・選択する完全な無細胞ワークフローを実現した点である。"
    "これにより、宿主域の制限を受けずに多様なファージゲノム変異体を迅速にスクリーニングできる基盤が確立された。"
    "従来 2〜4 週間を要していた宿主域改変を 1〜2 日で達成できる点は、ファージ療法の迅速なパーソナライズ化に直結する革新的意義を持つ。"
    "TXTL システムのスケーラビリティにより、高スループットな突然変異体スクリーニングへの発展も期待される。"
)

doc.add_heading("限界と今後の展望", level=2)
doc.add_paragraph(
    "現時点では T7 ファージ (40 kbp) を主なモデルとしており、より大型かつ複雑なゲノムを持つファージへの適用可能性は未検証である。"
    "TXTL 系のコスト低減と長期保存性の向上、さらに多様なファージ種への拡張が今後の課題であり、in vivo での宿主域変換効果の検証も必要である。"
)

doc.add_heading("トピックへのインプリケーション", level=2)
doc.add_paragraph(
    "本 PHEIGES プラットフォームはバクテリオファージの人工合成 (T2トピック) における核心技術であり、設計〜合成〜選択の全サイクルを無細胞化することで、合成ファージ工学の大幅な加速を実現する。"
)

today = datetime.date.today().isoformat()
doc.add_paragraph(f"Metadata: PMID=38472230; DOI=10.1038/s41467-024-46585-1; topic=2; generated={today}")

fname = safe_filename(title_en)
doc.save(fname)
print(f"Saved: {fname}")
