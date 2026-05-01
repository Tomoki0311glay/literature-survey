from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re, datetime

def set_yu_gothic(doc):
    style = doc.styles['Normal']
    style.font.name = 'Yu Gothic'
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), 'Yu Gothic')

def safe_filename(title, maxlen=150):
    s = re.sub(r'[<>:"/\\|?*]', '_', title)
    s = s.rstrip(' .')
    return s[:maxlen] + '.docx'

def add_table_row(table, label, value):
    row = table.add_row()
    row.cells[0].text = label
    row.cells[1].text = value

doc = Document()
set_yu_gothic(doc)

title_en = "Prediction of Klebsiella phage-host specificity at the strain level"
title_ja = "株レベルでのKlebsiellaファージ-宿主特異性の機械学習予測 (PhageHostLearn)"

doc.add_heading(title_en, level=1)

doc.add_heading("日本語訳タイトル", level=2)
doc.add_paragraph(title_ja)

doc.add_heading("論文情報", level=2)
table = doc.add_table(rows=1, cols=2)
table.style = 'Table Grid'
table.rows[0].cells[0].text = "項目"
table.rows[0].cells[1].text = "内容"
add_table_row(table, "著者", "Boeckaerts D, Stock M, Ferriol-González C, Oteo-Iglesias J, Sanjuán R, Domingo-Calap P, De Baets B, Briers Y")
add_table_row(table, "誌名", "Nature Communications")
add_table_row(table, "出版年", "2024")
add_table_row(table, "巻号頁", "15: 4355")
add_table_row(table, "DOI", "10.1038/s41467-024-48675-6")
add_table_row(table, "PMID", "38778023")

doc.add_heading("背景・課題", level=2)
doc.add_paragraph(
    "ファージ療法の成否は、ファージが標的となる細菌株に感染できるかどうかの精密なマッチングに依存する。"
    "従来の宿主域予測モデルは種レベルに留まっており、同一菌種内の株間で大きく異なる感染特異性を予測する手法が不足していた。"
    "特に Klebsiella pneumoniae (Kp) はカプセル型 (K型) や O型抗原の多様性により株特異性が高く、臨床ファージ選択における障壁となっていた。"
)

doc.add_heading("手法", level=2)
doc.add_paragraph(
    "PhageHostLearn は、ファージ受容体結合タンパク質 (RBP) と細菌の表面受容体 (K型カプセル多糖・O型抗原) の配列特徴を入力とするペアワイズ機械学習モデルである。"
    "Kp ファージ 72 株と宿主 Kp 菌株 178 株の感染実験データ (計 9,000 超のペア) を訓練・検証に使用した。"
    "RBP はファージゲノムから計算的に同定し、タンパク質言語モデル (ESM-1b) によるエンベディングで特徴量化した。"
    "宿主側は K型・O型の血清型情報をワンホットエンコードし、RBP エンベディングとの相互作用項をロジスティック回帰・勾配ブースティングで学習した。"
    "交差検証・外部バリデーション実験 (西スペイン由来の独立コホート) で性能を評価した。"
)

doc.add_heading("結果", level=2)
doc.add_paragraph(
    "内部交差検証において PhageHostLearn は ROC AUC 最大 81.8% を達成し、種レベル予測ベースラインを大幅に上回った。"
    "スペイン臨床株コホートでの外部バリデーション実験においても AUC 78.2% を維持し、汎化性能が示された。"
    "RBP のエンベディングが宿主特異性予測に最も寄与する特徴量であり、K型カプセルとの相互作用が主な決定因子であることが特徴量重要度解析で確認された。"
    "未知 K型を持つ菌株に対しても一定の予測性能を示し、ゼロショット予測の可能性を示唆した。"
    "モデルは計算コストが低く (推論 < 1 秒/ペア)、臨床現場での高スループットスクリーニングへの適用可能性が実証された。"
    "GitHub 上でオープンソースとして公開され、再現性と実用性が担保されている。"
)

doc.add_heading("考察・新規性", level=2)
doc.add_paragraph(
    "本研究の新規性は、ファージ-宿主相互作用予測を種レベルから株レベルへと初めて高精度で引き上げた点にある。"
    "RBP と細菌受容体の分子的相互作用を特徴量として明示的に組み込むことで、従来のゲノム類似性ベース予測では不可能だった株特異性の予測を実現した。"
    "臨床的な Kp 感染症に焦点を当てることで、即時の医療応用への道筋が明確であり、Nature Communications 掲載は当該分野への影響力を裏付ける。"
    "外部コホートでの検証は実装可能性を強化している。"
)

doc.add_heading("限界と今後の展望", level=2)
doc.add_paragraph(
    "現モデルは Klebsiella ファージに特化しており、他の菌種への直接適用は要再訓練。"
    "K型多糖の計算的同定精度に依存するため、新規 K型を持つ菌株の予測精度は低下する可能性があり、RBP 構造情報の統合や他菌種データセットへの拡張が今後の課題である。"
)

doc.add_heading("トピックへのインプリケーション", level=2)
doc.add_paragraph(
    "本研究はAI/MLによるファージ宿主予測 (T3トピック) の最前線であり、株レベルのマッチングという臨床上の最重要課題を機械学習で解決し、ファージ療法の実用化加速に直接貢献する。"
)

today = datetime.date.today().isoformat()
doc.add_paragraph(f"Metadata: PMID=38778023; DOI=10.1038/s41467-024-48675-6; topic=3; generated={today}")

fname = safe_filename(title_en)
doc.save(fname)
print(f"Saved: {fname}")
