from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re, datetime

def set_yu_gothic(doc):
    style = doc.styles['Normal']
    style.font.name = 'Yu Gothic'
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), 'Yu Gothic')

def safe_filename(title, maxlen=150):
    forbidden = r'[<>:"/\\|?*]'
    s = re.sub(forbidden, '_', title)
    s = s.rstrip(' .')
    return s[:maxlen] + '.docx'

def add_table_row(table, label, value):
    row = table.add_row()
    row.cells[0].text = label
    row.cells[1].text = value

doc = Document()
set_yu_gothic(doc)

title_en = "Large-scale phage cultivation for commensal human gut bacteria"
title_ja = "腸内共生細菌を対象とした大規模ファージ培養"

# Heading 1: English title
doc.add_heading(title_en, level=1)

# Heading 2 + paragraph: Japanese title
doc.add_heading("日本語訳タイトル", level=2)
doc.add_paragraph(title_ja)

# Heading 2 + table: metadata
doc.add_heading("論文情報", level=2)
table = doc.add_table(rows=1, cols=2)
table.style = 'Table Grid'
table.rows[0].cells[0].text = "項目"
table.rows[0].cells[1].text = "内容"
add_table_row(table, "著者", "Shen J, Zhang J, Mo L, Li Y, Li Y, Li C, Kuang X, Tao Z, Qu Z, Wu L, Chen J, Liu S, Zeng L, He Z, Chen Z, Deng Y, Zhang T, Li B, Dai L, Ma Y")
add_table_row(table, "誌名", "Cell Host & Microbe")
add_table_row(table, "出版年", "2023")
add_table_row(table, "巻号頁", "31(4): 665–677")
add_table_row(table, "DOI", "10.1016/j.chom.2023.03.013")
add_table_row(table, "PMID", "37054680")

# Background
doc.add_heading("背景・課題", level=2)
doc.add_paragraph(
    "腸内ファージはマイクロバイオームの構成を左右する重要な因子であるが、主要な腸内共生菌に感染するファージの系統的な単離・培養は進んでいなかった。"
    "既存のファージコレクションは病原菌や特定の菌種に偏っており、健康な腸内共生細菌を標的とするファージの網羅的なリソースが不足していた。"
    "この欠如がファージと宿主菌の相互作用解析や、精密菌叢改変療法の開発を阻んでいた。"
)

# Methods
doc.add_heading("手法", level=2)
doc.add_paragraph(
    "55 名の健常成人から採取した糞便検体を用い、42 菌種の腸内共生細菌に対して計 209 種のファージを単離した (Gut Phage Isolate Collection: GPIC)。"
    "単離は二重寒天プレーク法を基本とし、宿主菌の前培養条件・通気条件・フィルター細孔径を体系的に最適化することで、溶原性ファージの誘導も含めた広範な単離を実現した。"
    "ゲノム配列決定 (Illumina + Nanopore ハイブリッドアセンブリ) により全ゲノムを決定し、ファージの分類・生活環・溶菌特性を評価した。"
    "さらに Bacteroides fragilis を標的とする 8 ファージのカクテルを調製し、ヒト糞便由来の複合コミュニティ ex vivo モデルにて菌叢への効果を定量した。"
    "宿主域試験には代表株パネルを用い、ストレイン特異性を評価した。"
)

# Results
doc.add_heading("結果", level=2)
doc.add_paragraph(
    "42 菌種 (Bacteroides, Bifidobacterium, Faecalibacterium, Akkermansia 等の主要共生菌を含む) に対して計 209 種のファージが単離され、うち 70% 以上が新規ゲノムを有した。"
    "溶菌性ファージは全体の 30% 程度であり、残りは溶原性または慢性型であった。"
    "B. fragilis を対象とする 8 ファージカクテルは、健常ヒト糞便に由来する複合コミュニティ ex vivo モデルにおいて B. fragilis の相対存在量を最大 10 倍超減少させた。"
    "菌叢の他メンバーへの off-target 効果は有意には観察されず、カクテルの精密性が示された。"
    "ゲノム解析から、ファージの宿主特異性は主に尾部ファイバータンパク質の多様性に起因することが示された。"
    "本研究によりファージのライフスタイルや宿主域に関する大規模なデータセットが構築され、GPICはオープンリソースとして公開された。"
)

# Discussion
doc.add_heading("考察・新規性", level=2)
doc.add_paragraph(
    "本研究は腸内共生細菌を標的とする世界最大規模のファージコレクションを提供し、マイクロバイオーム研究の基盤リソースとなる。"
    "従来研究は病原菌ファージに偏っていたが、本研究は健常腸内共生菌ファージの系統的単離を初めて大規模に実現した点が新規性として高く評価される。"
    "複合コミュニティにおける B. fragilis の精密な減少は、疾患特異的な菌叢改変 (IBD、代謝疾患等) への応用可能性を具体的に示した。"
    "Cell Host & Microbe 誌掲載と多機関協力により、データの信頼性と再現性は高い。"
)

# Limitations
doc.add_heading("限界と今後の展望", level=2)
doc.add_paragraph(
    "現コレクションはヒト腸内 1,000 菌種超のうち 42 種に留まり、Firmicutes 等の偏性嫌気性菌ファージの単離が課題。"
    "In vivo での菌叢改変効果・安全性はマウスモデル以上の動物試験および臨床試験での検証が必要であり、また溶原性ファージの宿主内動態の予測モデル構築も今後の課題である。"
)

# Implications
doc.add_heading("トピックへのインプリケーション", level=2)
doc.add_paragraph(
    "本 GPIC コレクションはバクテリオファージによる腸内菌叢の精密改変 (T1トピック) の実現可能性を実証し、疾患ターゲット菌の選択的除去に向けた基盤ツールを提供する。"
)

# Metadata paragraph
today = datetime.date.today().isoformat()
doc.add_paragraph(f"Metadata: PMID=37054680; DOI=10.1016/j.chom.2023.03.013; topic=1; generated={today}")

fname = safe_filename(title_en)
doc.save(fname)
print(f"Saved: {fname}")
