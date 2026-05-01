from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re, datetime

def set_yu_gothic(doc):
    style = doc.styles['Normal']
    style.font.name = 'Yu Gothic'
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), 'Yu Gothic')

def safe_filename(title, maxlen=150):
    s = re.sub(r'[<>:"/\\|?*]', '_', title)
    s = s.rstrip(' .')
    return s[:maxlen] + '.docx'

def add_table_row(table, label, value):
    row = table.add_row()
    row.cells[0].text = label
    row.cells[1].text = value

doc = Document()
set_yu_gothic(doc)

title_en = "Molecular Mechanisms of Cationic Fusogenic Liposome Interactions with Bacterial Envelopes"
title_ja = "カチオン性融合リポソームと細菌エンベロープとの相互作用の分子メカニズム"

doc.add_heading(title_en, level=1)

doc.add_heading("日本語訳タイトル", level=2)
doc.add_paragraph(title_ja)

doc.add_heading("論文情報", level=2)
table = doc.add_table(rows=1, cols=2)
table.style = 'Table Grid'
table.rows[0].cells[0].text = "項目"
table.rows[0].cells[1].text = "内容"
add_table_row(table, "著者", "Scheeder A, Brockhoff M, Ward EN, Kaminski Schierle GS, Mela I, Kaminski CF")
add_table_row(table, "所属", "University of Cambridge (UK)")
add_table_row(table, "誌名", "Journal of the American Chemical Society (JACS)")
add_table_row(table, "出版年", "2024 (オンライン掲載 2023年12月)")
add_table_row(table, "巻号頁", "146(1): 1046–1057")
add_table_row(table, "DOI", "10.1021/jacs.3c11463")
add_table_row(table, "PMID", "38085801")

doc.add_heading("背景・課題", level=2)
doc.add_paragraph(
    "抗菌薬耐性菌の増加に伴い、細菌の外膜を越えて薬剤や核酸を直接送達する新規技術が求められている。"
    "リポソームは真核細胞への薬物送達で実績があるが、細菌 (特にグラム陰性菌の外膜二重構造) への効率的な輸送メカニズムは未解明な点が多く、合理的な送達システム設計の障壁となっていた。"
    "カチオン性融合リポソーム (CFL) が細菌エンベロープとどのように相互作用するかの分子レベルの理解が必要とされていた。"
)

doc.add_heading("手法", level=2)
doc.add_paragraph(
    "DOTAP/DOPE を主成分とするカチオン性融合リポソーム (直径 ~100 nm) を調製し、グラム陰性菌 Escherichia coli および グラム陽性菌 Bacillus subtilis との相互作用を超解像蛍光顕微鏡 (STED, 共焦点) で可視化した。"
    "蛍光ライフタイムイメージング顕微鏡 (FLIM) を用いてリポソーム膜と細菌膜の融合イベントをリアルタイムで定量した。"
    "脂質移行 (lipid mixing) アッセイと内容物漏出 (content leakage) アッセイを組み合わせて融合の真正性を確認した。"
    "バンコマイシン (通常グラム陰性菌外膜を透過できない糖ペプチド抗生物質) を CFL に封入し、E. coli に対する共送達効果を最小発育阻止濃度 (MIC) および生菌数測定で評価した。"
    "共焦点ライブイメージングにより、リポソームの細菌表面への結合・融合・内容物放出の動態を経時的に解析した。"
)

doc.add_heading("結果", level=2)
doc.add_paragraph(
    "E. coli (グラム陰性菌) において、CFL は外膜に直接融合し、人工脂質を外膜に統合することで膜の流動性を変化させ、細菌の生存率を低下させた。"
    "STED 超解像顕微鏡により、CFL が E. coli 外膜上で局所的な膜変形および孔形成を誘導することが可視化された。"
    "B. subtilis (グラム陽性菌) では融合ではなく付着・脂質インターナリゼーションが主な相互作用形式であることが FLIM で示された。"
    "バンコマイシンを封入した CFL は、通常では E. coli に無効な濃度 (MIC > 256 µg/ml) においても感染性を有意に低下させ、外膜を迂回した直接ペリプラズム送達が示唆された。"
    "CFL の融合効率は細菌膜のカルジオリピン含量と正の相関を示し、膜組成が送達効率の決定因子であることが示された。"
    "融合は 10 分以内に開始し、30 分以内に最大値に達する高速プロセスであった。"
)

doc.add_heading("考察・新規性", level=2)
doc.add_paragraph(
    "本研究の核心的新規性は、超解像顕微鏡と FLIM を組み合わせてリポソーム-細菌融合を分子レベルで初めて詳細に可視化した点にある。"
    "グラム陰性菌とグラム陽性菌で融合メカニズムが根本的に異なることを実証し、設計ガイドラインを提供した。"
    "バンコマイシンの E. coli ペリプラズム送達成功は、「抗菌スペクトルの壁」を CFL で乗り越える概念実証として産業・医療両面で高いインパクトを持つ。"
)

doc.add_heading("限界と今後の展望", level=2)
doc.add_paragraph(
    "本研究は主に in vitro モデルに限定されており、生体内 (腸管粘液層、血清タンパク質との非特異的相互作用) での融合効率は別途検証が必要である。"
    "また、脂質組成の最適化・ターゲティングリガンド付加による菌種選択性の向上、および核酸 (プラスミド・アンチセンス RNA) 送達への応用が今後の重要課題である。"
)

doc.add_heading("トピックへのインプリケーション", level=2)
doc.add_paragraph(
    "本研究はリポソームによる細菌への核酸/タンパク質輸送 (T4トピック) の物理化学的基盤を提供し、グラム陰性菌外膜という最大の障壁を融合メカニズムで克服する経路を分子レベルで実証する。"
)

today = datetime.date.today().isoformat()
doc.add_paragraph(f"Metadata: PMID=38085801; DOI=10.1021/jacs.3c11463; topic=4; generated={today}")

fname = safe_filename(title_en)
doc.save(fname)
print(f"Saved: {fname}")
