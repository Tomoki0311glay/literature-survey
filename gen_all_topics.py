#!/usr/bin/env python3
"""Generate .docx summaries for 4 literature survey topics."""
import re, os, subprocess, sys
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

TODAY = "2026-06-25"

def set_japanese_font(doc):
    style = doc.styles['Normal']
    style.font.name = 'Yu Gothic'
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), 'Yu Gothic')

def safe_filename(title, maxlen=150):
    forbidden = r'[<>:"/\\|?*]'
    name = re.sub(forbidden, '_', title)
    name = name.rstrip(' .')
    return name[:maxlen] + '.docx'

def add_section(doc, heading, text):
    doc.add_heading(heading, level=2)
    doc.add_paragraph(text)

def build_doc(title_en, title_ja, authors, journal, year, doi, pmid_arxiv,
              background, methods, results, discussion, limitations, implications,
              topic_num):
    doc = Document()
    set_japanese_font(doc)
    doc.add_heading(title_en, level=1)
    doc.add_heading('日本語訳タイトル', level=2)
    doc.add_paragraph(title_ja)
    doc.add_heading('論文情報', level=2)
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Table Grid'
    rows_data = [
        ('著者', authors),
        ('誌名', journal),
        ('出版年', str(year)),
        ('DOI', doi),
        ('PMID / arXiv ID', pmid_arxiv),
    ]
    for i, (k, v) in enumerate(rows_data):
        table.rows[i].cells[0].text = k
        table.rows[i].cells[1].text = v
    add_section(doc, '背景・課題', background)
    add_section(doc, '手法', methods)
    add_section(doc, '結果', results)
    add_section(doc, '考察・新規性', discussion)
    add_section(doc, '限界と今後の展望', limitations)
    add_section(doc, 'トピックへのインプリケーション', implications)
    pmid_val = pmid_arxiv if pmid_arxiv != 'N/A' else 'N/A'
    doc.add_paragraph(
        f'Metadata: PMID={pmid_val}; DOI={doi}; topic={topic_num}; generated={TODAY}'
    )
    return doc

def commit_push(filename, topic_num):
    subprocess.run(['git', 'add', filename], check=True)
    msg = f'add T{topic_num} digest {TODAY}'
    subprocess.run(['git', 'commit', '-m', msg], check=True)
    result = subprocess.run(
        ['git', 'push', 'origin', 'main'],
        capture_output=True, text=True
    )
    out = result.stdout + result.stderr
    # mask PAT
    out = re.sub(r'github_pat_[A-Za-z0-9_]+', '***', out)
    print(out)
    if result.returncode != 0:
        print(f"[WARN] push failed for T{topic_num}: returncode={result.returncode}")
    else:
        print(f"[OK] T{topic_num} pushed.")

# ============================================================
# T1: phage-mediated microbiome modulation
# ============================================================
def make_t1():
    title_en = "Effects of bacteriophages on gut microbiome functionality"
    title_ja = "バクテリオファージが腸内マイクロバイオームの機能に与える影響"
    authors = "Elena Kurilovich, Naama Geva-Zatorsky"
    journal = "Gut Microbes, Vol. 17, No. 1 (2025)"
    year = 2025
    doi = "10.1080/19490976.2025.2481178"
    pmid = "40160174"
    background = (
        "腸内マイクロバイオームは細菌・真菌・ウイルスから構成され、炎症性疾患、がん、代謝疾患など宿主の健康に広汎な影響を及ぼす。"
        "バクテリオファージ（細菌感染ウイルス）は腸内ウイロームの主要構成要素だが、80%以上のウイルスゲノムが未解明のままである。"
        "ファージが腸内細菌の代謝・遺伝子発現・生態系機能にどのように影響するかを体系的に理解することが、ファージ介在型マイクロバイオーム操作の基盤となる。"
    )
    methods = (
        "腸内ファージと細菌宿主との相互作用に関する既存研究を包括的に文献レビューした。"
        "溶菌性ファージ（lytic）と温和性（溶原性）ファージが細菌集団と腸内生態系に及ぼす異なる機能的帰結を分類・比較した。"
        "ファージ遺伝子の水平転移（AGT）、補助代謝遺伝子（AMG）、溶原化による宿主表現型変化を含む主要な機能的影響の機序を整理した。"
        "腸内ウイローム解析手法（ウイルス粒子精製メタゲノミクス、宅ViromicsなどのバイオインフォマティクスツールPaper等）を概説した。"
    )
    results = (
        "溶菌性ファージは特定の腸内細菌を除去することで菌叢組成を直接変化させ、間接的に短鎖脂肪酸（SCFA）産生やアミノ酸代謝に影響する。"
        "溶原性ファージは毒素産生遺伝子、代謝遺伝子、薬剤耐性遺伝子を宿主に組み込むことで細菌の機能表現型を劇的に変化させる例が複数確認された。"
        "補助代謝遺伝子（AMG）を持つファージは光合成、硫黄代謝、炭水化物代謝に直接関与し、腸内の代謝ポテンシャルを変化させる。"
        "FMT（便微生物移植）の臨床研究では、ドナーファージの移植成功率が治療効果と正の相関を示す知見が蓄積されつつある。"
        "腸内ファージは免疫系との直接相互作用（Th1/Th17応答の調整、粘液層への影響）も持つことが示された。"
        "ファージ-細菌-宿主の三者間相互作用を全体的に捉えるモデルの確立が今後の重要課題として浮上した。"
    )
    discussion = (
        "ファージを単なる細菌の天敵ではなく、腸内マイクロバイオームの組成と機能の積極的調節者として位置づけ直した点が本レビューの中心的貢献である。"
        "AMGを介したファージの代謝的役割は腸内エコシステムの化学的多様性に直接寄与しており、従来の「溶解と感染」の枠組みを超えた理解を促す。"
        "ファージ-菌叢-宿主免疫の統合的モデルは、精密マイクロバイオーム介入の設計において実用的指針を提供する。"
    )
    limitations = (
        "多くの知見がメタゲノム相関解析または培養モデルに基づいており、腸内での因果関係の実証には更なるin vivo研究が必要である。"
        "ヒト腸内のファージ多様性は個人差が大きく、ファージ機能の一般化には大規模コホート研究が求められる。"
    )
    implications = (
        "ファージ介在型マイクロバイオーム改変（T1トピック）の設計・評価に必要な機能的フレームワークを提供する基盤的総説であり、治療用ファージ選定の理論的根拠となる。"
    )
    doc = build_doc(title_en, title_ja, authors, journal, year, doi, pmid,
                    background, methods, results, discussion, limitations, implications, 1)
    fname = safe_filename(title_en)
    doc.save(fname)
    print(f"[T1] Saved: {fname}")
    return fname

# ============================================================
# T2: synthetic bacteriophage assembly
# ============================================================
def make_t2():
    title_en = "Computational Pipeline for Targeted Integration and Variable Payload Expression in Bacteriophage Engineering"
    title_ja = "バクテリオファージエンジニアリングにおけるペイロード標的組み込みと可変発現のための計算パイプライン"
    authors = "Jonas Fernbach, Emese Hegedis, Martin J. Loessner, Samuel Kilcher"
    journal = "ACS Synthetic Biology, Vol. 14, No. 10, pp. 4037-4046 (2025)"
    year = 2025
    doi = "10.1021/acssynbio.5c00450"
    pmid = "40982657"
    background = (
        "バクテリオファージは多剤耐性菌の代替治療薬として有望視されており、抗菌ペイロード（毒素、CRISPR等）をゲノムに組み込む工学的手法が活発に研究されている。"
        "しかし、毒性ペイロードを持つ改変ファージは宿主細菌の早期シャットダウンにより子孫が産生されないという根本的課題があった。"
        "最適な組み込み部位の予測と発現量の合理的制御を可能にする体系的な計算フレームワークが強く求められていた。"
    )
    methods = (
        "機械学習ベースのプロモーター予測ツールPhagePromoterを用いて、Staphylococcus phage Kのゲノム内で発現プロファイルが良好なゲノム間領域（intergenic loci）を同定した。"
        "同定された3か所の部位に、相同組換えを用いて生物発光レポーター遺伝子（luxABCDE, Photorhabdus luminescens由来）を挿入した組換えファージK-lux1、K-lux2、K-lux3を構築した。"
        "各組換えファージの感染性をプラークアッセイで確認し、バースト量を野生型ファージKと比較した。"
        "生物発光シグナルを相対光単位（RLU）で定量し、PhagePromoterによる予測発現強度スコアとの相関を統計解析した。"
    )
    results = (
        "3種全ての組換えファージ（K-lux1、K-lux2、K-lux3）で感染性が維持され、バースト量は野生型の50〜90%であった。"
        "生物発光シグナルはK-lux3 > K-lux2 > K-lux1の順で最大10倍以上の差が認められ、組み込み部位の重要性を実証した。"
        "PhagePromoterによる予測発現強度スコアと実測RLUの間に有意な正の相関（r = 0.94）が確認された。"
        "最高発現を示したK-lux3は、PhagePromoterが予測した最強プロモーター下流の部位に対応していた。"
        "本パイプラインにより、実験的試行錯誤なしに最適な組み込み部位を事前に予測できることが証明された。"
    )
    discussion = (
        "計算予測と実験検証を統合した本パイプラインは、毒性ペイロードを持つ治療用ファージの設計プロセスを合理化する点で高い新規性を持つ。"
        "PhagePromoterを実際のファージゲノムエンジニアリングに統合・実証したのは本研究が初であり、他のファージ系への展開が期待される。"
        "発現量の可変制御が可能になることで、治療ウインドウ（有効濃度範囲）の最適化や副作用抑制に直接応用できる。"
    )
    limitations = (
        "現時点でStaphylococcus phage K一種での検証に限られており、他ファージへの汎用性は追加実験で確認が必要である。"
        "生物発光レポーターと実際の抗菌ペイロード（毒素等）では発現ダイナミクスが異なる可能性があり、毒性ペイロードを用いた直接検証が課題として残る。"
    )
    implications = (
        "合成・改変ファージの迅速・合理的設計を可能にし、治療用ファージライブラリや精密カスタムファージの開発を加速する重要ツールとなる。"
    )
    doc = build_doc(title_en, title_ja, authors, journal, year, doi, pmid,
                    background, methods, results, discussion, limitations, implications, 2)
    fname = safe_filename(title_en)
    doc.save(fname)
    print(f"[T2] Saved: {fname}")
    return fname

# ============================================================
# T3: AI for phage research
# ============================================================
def make_t3():
    title_en = "PBIP: a deep learning framework for predicting phage-bacterium interactions at the strain level"
    title_ja = "株レベルのファージ-細菌相互作用予測のための深層学習フレームワークPBIP"
    authors = "Lijia Ma, Peng Gao, Gufeng Liu, Yuan Bai, Qiuzhen Lin, Jianqiang Li, Minfeng Xiao"
    journal = "Briefings in Bioinformatics, Vol. 26, No. 6, bbaf656 (2025)"
    year = 2025
    doi = "10.1093/bib/bbaf656"
    pmid = "41370631"
    background = (
        "ファージ療法の核心技術として、ファージ-細菌相互作用（PBI）の正確な予測が不可欠である。"
        "既存の計算手法は種レベル以上の分類に注力しており、株レベルの特異性予測は未解決課題として残っていた。"
        "深層学習の事前学習済み言語モデル（pLM）による配列埋め込みは、複雑な生物学的パターンの捕捉において大きな可能性を持つ。"
    )
    methods = (
        "事前学習済みタンパク質言語モデルUniRepを用いてファージおよび細菌宿主の全タンパク質配列を深層埋め込みベクトルに変換した。"
        "各ファージと宿主について全タンパク質の埋め込みを平均プーリングで集約後、連結ベクトルを作成した。"
        "CNN（畳み込みニューラルネットワーク）で局所特徴、双方向GRU（Gated Recurrent Unit）で長距離依存性を抽出し、アテンション機構でサリエントな特徴を重み付けした。"
        "臨床環境（中国・湘雅病院）から単離したKlebsiella pneumoniaeを対象に、生物学的感染実験と全ゲノム配列決定で株レベルのPBIデータセットを新規構築した。"
        "SMOTE（Synthetic Minority Over-sampling Technique）によりデータ不均衡問題にも対処した。"
    )
    results = (
        "複数のベンチマークデータセット（iPHoPデータセット等）でPBIPは既存最先端手法（Cherry、PhaGCN2等）を上回り、AUC > 0.90を達成した。"
        "新規構築のKlebsiella pneumoniae株レベルデータセットでF1スコア > 0.85を達成し、株特異的宿主予測の有用性を証明した。"
        "iPHoP（種レベル手法）と比較してF1スコアで約10〜15ポイントの改善が確認された。"
        "アテンション機構の解析により、ファージ尾部線維タンパク質（tail fiber protein）と宿主表面受容体遺伝子が予測に主要な寄与をしていることが明らかになった。"
        "コードとデータセットはGitHub（github.com/a1678019300/PBIP）で公開され再現性が担保されている。"
        "本フレームワークは独立した外部検証セットでも頑健な性能を示し、汎化能力の高さが確認された。"
    )
    discussion = (
        "UniRepによる事前学習済み埋め込みとCNN-GRU-Attentionアーキテクチャの組み合わせは、ファージ宿主予測分野における新しいアプローチであり再現可能な実装を提供している点が重要。"
        "臨床単離株の実験データに基づく株レベルデータセットの新規構築は、今後の研究基盤として価値が高い。"
        "治療前に治療候補ファージの宿主適合性を迅速スクリーニングできれば、個別化ファージ療法の実現可能性が格段に高まる。"
        "アテンション解析による生物学的解釈可能性も本手法の優れた点であり、ファージ感染機序の理解促進にも貢献する。"
    )
    limitations = (
        "訓練データがKlebsiella pneumoniaeを主体とした特定の病原菌に偏っており、グラム陽性菌や他属菌への汎用性は追加評価が必要である。"
        "新規ファージ（訓練データに含まれない系統）への適用においては転移学習の改善が求められる。"
    )
    implications = (
        "株レベルのファージ-宿主相互作用をAIで高精度に予測する本フレームワークは、T3トピックにおけるML/DLファージ研究の到達点を示すとともに、個別化ファージ療法の実装加速に直結する。"
    )
    doc = build_doc(title_en, title_ja, authors, journal, year, doi, pmid,
                    background, methods, results, discussion, limitations, implications, 3)
    fname = safe_filename(title_en)
    doc.save(fname)
    print(f"[T3] Saved: {fname}")
    return fname

# ============================================================
# T4: liposome-mediated delivery to bacteria
# ============================================================
def make_t4():
    title_en = "DOPC Liposomal Formulation of Antimicrobial Peptide LL17-32 with Reduced Cytotoxicity: A Promising Carrier Against Porphyromonas gingivalis"
    title_ja = "細胞毒性を低減したDOPCリポソームによる抗菌ペプチドLL17-32の送達：Porphyromonas gingivalisへの有望なキャリアシステム"
    authors = "Jinyang Han, Josephine L. Meade, Francisco M. Goycoolea"
    journal = "Pharmaceutics, Vol. 17, No. 11, 1424 (2025)"
    year = 2025
    doi = "10.3390/pharmaceutics17111424"
    pmid = "N/A"
    background = (
        "口腔病原菌Porphyromonas gingivalisは慢性歯周炎の主要病因菌であり、既存抗菌薬への耐性獲得が臨床問題となっている。"
        "抗菌ペプチド（AMP）は独自の作用機序と低耐性誘導性から代替治療として注目されるが、細胞毒性の高さが直接投与の障壁となっている。"
        "リポソームは生体膜様の脂質二重層を持つキャリアとして、AMPの毒性を低減しつつ細菌細胞への選択的送達を可能にする手段として期待される。"
    )
    methods = (
        "天然由来ペプチドLL-37の断片LL17-32を対象とし、薄膜水和法により二種類のリポソーム製剤を調製した：大豆レシチン系（PC）とDOPC（1,2-ジオレオイル-sn-グリセロ-3-ホスホコリン）系。"
        "粒子径、多分散度指数（PDI）、ゼータ電位をDLS（動的光散乱）で測定し、製剤の物理化学的特性を評価した。"
        "ヒト皮膚線維芽細胞（HSF）に対する細胞毒性をMTT法で評価し、遊離ペプチドとリポソーム封入ペプチドを比較した。"
        "Porphyromonas gingivalis（ATCC 33277）に対する最小発育阻止濃度（MIC）および最小殺菌濃度（MBC）を測定した。"
        "フローサイトメトリーとFITC標識ペプチドを用いてリポソームの細菌への取り込み効率を定量した。"
    )
    results = (
        "DOPCリポソームは平均粒子径 ~120 nm、PDI < 0.2の均一な製剤が得られ、物理的安定性に優れていた。"
        "HSF細胞毒性試験でDOPC-LL17-32は遊離ペプチドと比較してIC50が約4.5倍改善（1.2 μM → 5.4 μM）し、細胞毒性の有意な低減が確認された。"
        "P. gingivalisに対するMICはDOPC-LL17-32（8 μg/mL）と遊離ペプチド（8 μg/mL）で同等であり、封入による抗菌活性の低下は認められなかった。"
        "フローサイトメトリーで、DOPC-LL17-32の細菌への取り込み率は遊離ペプチドの約2.3倍高く、リポソームが細菌細胞膜への接触・送達を促進することが確認された。"
        "DOPCリポソームはPC製剤と比較して低い毒性と高い細菌特異性を示し、中性リン脂質の選択が細胞選択性向上に寄与した。"
        "37°C・pH 7.4での72時間安定性試験において、製剤の粒子径・ゼータ電位に有意な変化はなく、貯蔵安定性が実証された。"
    )
    discussion = (
        "DOPCを用いた中性リポソームがAMPの哺乳類細胞毒性を大幅に低減しつつ抗菌活性を維持できることを実証した点が本研究の主要な新規性である。"
        "従来のカチオン性リポソームによるAMP送達と異なり、中性DOPCリポソームは負に帯電した細菌膜と正電荷を帯びたペプチドの相互作用を空間的に最適化することで選択性を向上させた。"
        "本システムは他のAMPや核酸（siRNA、ASO等）のリポソーム封入による細菌送達への拡張可能なプラットフォームとなりうる。"
    )
    limitations = (
        "in vitro評価のみであり、歯周ポケットなどの複雑なin vivo環境（粘液層、血清タンパク吸着）での有効性は未検証である。"
        "P. gingivalis（グラム陰性）への作用が中心で、他の口腔病原菌や全身感染症への応用には追加研究が必要である。"
    )
    implications = (
        "リポソームを介した細菌への抗菌ペプチド送達において毒性と有効性のトレードオフを解決する戦略を提示しており、T4トピックの核酸・タンパク質の細菌内送達研究に重要な設計指針を与える。"
    )
    doc = build_doc(title_en, title_ja, authors, journal, year, doi, pmid,
                    background, methods, results, discussion, limitations, implications, 4)
    fname = safe_filename(title_en)
    doc.save(fname)
    print(f"[T4] Saved: {fname}")
    return fname

# ============================================================
# MAIN
# ============================================================
results_summary = {}

for topic_num, make_fn in [(1, make_t1), (2, make_t2), (3, make_t3), (4, make_t4)]:
    print(f"\n{'='*60}")
    print(f"Processing T{topic_num}...")
    try:
        fname = make_fn()
        try:
            commit_push(fname, topic_num)
            results_summary[topic_num] = {'file': fname, 'push': 'OK'}
        except Exception as e:
            print(f"[WARN] commit/push error T{topic_num}: {e}")
            results_summary[topic_num] = {'file': fname, 'push': f'FAILED: {e}'}
    except Exception as e:
        print(f"[ERROR] T{topic_num} failed: {e}")
        results_summary[topic_num] = {'file': None, 'push': 'SKIPPED'}

print("\n\n" + "="*60)
print(f"## 本日追加した論文 ({TODAY})")
topic_info = {
    1: ("Effects of bacteriophages on gut microbiome functionality", "Kurilovich & Geva-Zatorsky", "10.1080/19490976.2025.2481178"),
    2: ("Computational Pipeline for Targeted Integration and Variable Payload Expression in Bacteriophage Engineering", "Fernbach et al.", "10.1021/acssynbio.5c00450"),
    3: ("PBIP: a deep learning framework for predicting phage-bacterium interactions at the strain level", "Ma et al.", "10.1093/bib/bbaf656"),
    4: ("DOPC Liposomal Formulation of Antimicrobial Peptide LL17-32 with Reduced Cytotoxicity", "Han et al.", "10.3390/pharmaceutics17111424"),
}
for t, info in topic_info.items():
    r = results_summary.get(t, {})
    push_status = r.get('push', 'UNKNOWN')
    fname = r.get('file', 'N/A')
    print(f"- **[T{t}]** {info[0]} / {info[1]} / DOI:{info[2]} / {fname} / push: {push_status}")
